import pytest
import os
import docx
import openpyxl
from exam.utils import parse_questions, parse_student_excel

@pytest.fixture
def sample_txt_file(tmp_path):
    content = """What is the capital of France?
a) London
b) Paris
c) Berlin
d) Madrid
Correct: (b)

Problem Statement: Write a function to find the factorial of a number.
The function should handle both positive integers and edge cases.

What is 2 + 2?
a) 3
b) 4
c) 5
d) 6
Correct: (b)"""
    
    file_path = tmp_path / "test_questions.txt"
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    return str(file_path)

@pytest.fixture
def sample_docx_file(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("What is the capital of France?")
    doc.add_paragraph("a) London")
    doc.add_paragraph("b) Paris")
    doc.add_paragraph("c) Berlin")
    doc.add_paragraph("d) Madrid")
    doc.add_paragraph("Correct: (b)")
    doc.add_paragraph("")
    doc.add_paragraph("Problem Statement: Write a function to find the factorial of a number.")
    doc.add_paragraph("The function should handle both positive integers and edge cases.")
    
    file_path = tmp_path / "test_questions.docx"
    doc.save(file_path)
    return str(file_path)

@pytest.fixture
def sample_excel_file(tmp_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    
    # Add headers
    ws.append(["Name", "Username", "Email"])
    
    # Add sample data
    ws.append(["John Doe", "johndoe", "john@example.com"])
    ws.append(["Jane Smith", "janesmith", "jane@example.com"])
    ws.append(["Bob Johnson", "bobjohnson", "bob@example.com"])
    
    file_path = tmp_path / "test_students.xlsx"
    wb.save(file_path)
    return str(file_path)

class TestParseQuestions:
    def test_parse_questions_txt(self, sample_txt_file):
        result = parse_questions(sample_txt_file, "test_questions.txt")
        
        # Verify structure
        assert "examId" in result
        assert "questions" in result
        assert "maxScore" in result
        
        # Verify questions
        questions = result["questions"]
        assert len(questions) == 3  # 2 MCQ + 1 coding question
        
        # Verify MCQ questions
        mcq_questions = [q for q in questions if q["type"] == "mcq"]
        assert len(mcq_questions) == 2
        
        # Verify first MCQ
        first_mcq = mcq_questions[0]
        assert "What is the capital of France?" in first_mcq["question"]
        assert len(first_mcq["options"]) == 4
        assert first_mcq["correctAnswer"] == "b"
        assert first_mcq["score"] == 2
        
        # Verify coding question
        coding_questions = [q for q in questions if q["type"] == "coding"]
        assert len(coding_questions) == 1
        coding_q = coding_questions[0]
        assert "factorial" in coding_q["question"].lower()
        assert coding_q["score"] == 5
        
        # Verify total score
        assert result["maxScore"] == 9  # 2 MCQ (2 points each) + 1 coding (5 points)

    def test_parse_questions_docx(self, sample_docx_file):
        result = parse_questions(sample_docx_file, "test_questions.docx")
        
        # Verify structure
        assert "examId" in result
        assert "questions" in result
        assert "maxScore" in result
        
        # Verify questions
        questions = result["questions"]
        assert len(questions) == 2  # 1 MCQ + 1 coding question
        
        # Verify MCQ question
        mcq_questions = [q for q in questions if q["type"] == "mcq"]
        assert len(mcq_questions) == 1
        assert "capital of France" in mcq_questions[0]["question"]
        
        # Verify coding question
        coding_questions = [q for q in questions if q["type"] == "coding"]
        assert len(coding_questions) == 1
        assert "factorial" in coding_questions[0]["question"].lower()

    def test_parse_questions_with_custom_uid(self, sample_txt_file):
        custom_uid = "test-exam-123"
        result = parse_questions(sample_txt_file, "test_questions.txt", custom_uid)
        assert result["examId"] == custom_uid

    def test_parse_questions_empty_file(self, tmp_path):
        # Create empty file
        empty_file = tmp_path / "empty.txt"
        empty_file.write_text("")
        
        result = parse_questions(str(empty_file), "empty.txt")
        assert len(result["questions"]) == 0
        assert result["maxScore"] == 0

class TestParseStudentExcel:
    def test_parse_student_excel_valid_data(self, sample_excel_file):
        students = parse_student_excel(sample_excel_file)
        
        # Verify number of students
        assert len(students) == 3
        
        # Verify student data structure
        for student in students:
            assert "name" in student
            assert "username" in student
            assert "email" in student
        
        # Verify specific student data
        assert students[0]["name"] == "John Doe"
        assert students[0]["username"] == "johndoe"
        assert students[0]["email"] == "john@example.com"
        
        assert students[1]["name"] == "Jane Smith"
        assert students[1]["username"] == "janesmith"
        assert students[1]["email"] == "jane@example.com"

    def test_parse_student_excel_empty_file(self, tmp_path):
        # Create empty Excel file
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Name", "Username", "Email"])  # Only headers
        
        empty_file = tmp_path / "empty_students.xlsx"
        wb.save(empty_file)
        
        students = parse_student_excel(str(empty_file))
        assert len(students) == 0

    def test_parse_student_excel_incomplete_data(self, tmp_path):
        # Create Excel file with incomplete data
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Name", "Username", "Email"])
        ws.append(["John Doe", "", "john@example.com"])  # Missing username
        ws.append(["", "janesmith", "jane@example.com"])  # Missing name
        ws.append(["Bob Johnson", "bobjohnson", ""])  # Missing email
        
        incomplete_file = tmp_path / "incomplete_students.xlsx"
        wb.save(incomplete_file)
        
        students = parse_student_excel(str(incomplete_file))
        assert len(students) == 0  # Should skip rows with missing data 